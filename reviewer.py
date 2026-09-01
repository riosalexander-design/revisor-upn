import os
import io
import fitz  # PyMuPDF
import docx
import re
import json
import time
from google import genai
from google.genai import types

import pandas as pd

def extract_text_from_uploaded_file(uploaded_file):
    name = uploaded_file.name.lower()
    text = ""
    try:
        if name.endswith('.pdf'):
            pdf_bytes = uploaded_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text()
        elif name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif name.endswith('.txt'):
            text = uploaded_file.read().decode('utf-8', errors='ignore')
        elif name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
            text = df.to_string()
        elif name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
            text = df.to_string()
    except Exception as e:
        print(f"Error reading {name}: {e}")
    return text

def extract_text_from_local_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
        elif ext == '.docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    return text

def evaluate_project(uploaded_files, normative_folder, api_key):
    client = genai.Client(api_key=api_key)
    
    # Extraer texto de los archivos subidos en Streamlit
    document_text = ""
    for uf in uploaded_files:
        document_text += f"\n\n--- INICIO DE DOCUMENTO: {uf.name} ---\n\n"
        document_text += extract_text_from_uploaded_file(uf)
        document_text += f"\n\n--- FIN DE DOCUMENTO: {uf.name} ---\n\n"
        
    if not document_text.strip():
        return "No se encontraron documentos válidos o están vacíos."

    # Extraer texto de la carpeta de normativas local del servidor
    upn_normative = ""
    if os.path.isdir(normative_folder):
        for root, dirs, files in os.walk(normative_folder):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx', '.txt')):
                    path = os.path.join(root, file)
                    upn_normative += f"\n\n--- NORMATIVA: {file} ---\n\n"
                    upn_normative += extract_text_from_local_file(path)
    else:
        return f"No se encontró la carpeta de normativas en el servidor: {normative_folder}"
    
    sys_instruction = (
        "Eres un panel de 5 expertos (Asesor Científico de Nivel Doctoral, Metodólogo, Verificador Bibliográfico, Corrector de Estilo Vancouver y Revisor Normativo) y un Editor/Calificador Final.\n"
        "Tu tarea es generar un INFORME DE REVISIÓN DEL PROYECTO DE INVESTIGACIÓN unificado, estructurado de la siguiente manera:\n\n"
        "1. REPORTE CIENTÍFICO-DOCTORAL: Evalúa la solidez y coherencia lógica del planteamiento del problema, los objetivos, y la relevancia científica general de la investigación.\n"
        "2. REPORTE METODOLÓGICO: Identifica debilidades metodológicas, de diseño y estadísticas.\n"
        "3. REPORTE BIBLIOGRÁFICO: Verifica la relevancia científica de fuentes (priorizando PubMed, SciELO, etc.).\n"
        "4. REPORTE ESTILO VANCOUVER: Revisa la alineación estricta a las normas Vancouver.\n"
        "5. REPORTE NORMATIVA INSTITUCIONAL: Verifica cumplimiento de formatos y normativa institucional.\n\n"
        "Para CADA debilidad encontrada en estas 5 secciones, mantén ESTRICTAMENTE este formato:\n"
        "OBSERVACIÓN: [Descripción puntual y precisa del error]\n"
        "EXPLICACIÓN: [Motivos con sustento]\n"
        "SUGERENCIA DE CORRECCIÓN: [Instrucción imperativa para el estudiante]\n\n"
        "Al final de todo el informe, debes entregar una CALIFICACIÓN FINAL aplicando la rúbrica proporcionada, restando puntos según las observaciones, y explicando brevemente cómo obtuviste esa nota.\n\n"
        "Tu tarea final es MUY IMPORTANTE: Al puro final de tu respuesta, debes agregar un bloque JSON estricto con la siguiente estructura (y nada más después):\n"
        "```json\n"
        "{\n"
        '  "score": [tu calificación final en formato numérico],\n'
        '  "has_ethics": [true o false, si el proyecto ya cuenta con evaluación o aprobación del comité de ética],\n'
        '  "has_grave_obs_before_methodology": [true o false, si tiene observaciones graves o muy serias hasta la metodología]\n'
        "}\n"
        "```\n\n"
        "IMPORTANTE: Redacta absolutamente todo el informe (fuera del bloque JSON) en TEXTO PLANO y EN SEGUNDA PERSONA. "
        "Háblale directamente al estudiante o autor del proyecto (ejemplo: 'Tu planteamiento del problema es...', 'Debes corregir...', 'Has citado incorrectamente...'). "
        "NO uses formato Markdown. NO uses asteriscos (**) para negritas, NO uses numerales (#) para títulos, ni guiones (-) para listas. "
        "Usa solo letras mayúsculas para los títulos y números normales si necesitas listar algo."
    )

    prompt = (
        f"Analiza los siguientes documentos del proyecto de investigación:\n\n{document_text}\n\n"
        f"Considera la siguiente normativa/rúbrica de la universidad:\n{upn_normative}"
    )

    models_to_try = [
        'gemini-3.6-flash',
        'gemini-3.5-flash',
        'gemini-3.7-flash',
        'gemini-flash-latest'
    ]
    
    last_error = ""

    for model_name in models_to_try:
        for attempt in range(3): # Reintenta 3 veces por cada modelo
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=sys_instruction,
                        temperature=0.2,
                    )
                )
                
                raw_report = response.text
                
                # Corrección de Mojibake (Si la API devuelve texto mal decodificado en ISO-8859-1 en lugar de UTF-8)
                try:
                    raw_report = raw_report.encode('latin-1').decode('utf-8')
                except UnicodeError:
                    pass
                
                
                metadata = {
                    "score": 0,
                    "has_ethics": False,
                    "has_grave_obs_before_methodology": True
                }
                report_clean = raw_report
                
                json_match = re.search(r'```json\s*(\{.*?\})\s*```', raw_report, re.DOTALL)
                if json_match:
                    try:
                        metadata = json.loads(json_match.group(1))
                        report_clean = raw_report[:json_match.start()].strip()
                    except:
                        pass

                return {
                    "report": report_clean,
                    "metadata": metadata
                }
                
            except Exception as e:
                last_error = str(e)
                # Si el modelo no existe o está bloqueado para esta API Key (404), no perdemos tiempo reintentando, pasamos al siguiente modelo.
                if "404" in last_error or "NOT_FOUND" in last_error or "is no longer available" in last_error:
                    break 
                
                # Si es un error de saturación temporal (503), esperamos un poco y reintentamos con el mismo modelo
                time.sleep(5 * (attempt + 1))
                
    return {"report": f"Error Crítico: Servidores de IA saturados. Se intentó conectar con todos los modelos disponibles sin éxito. Último error: {last_error}", "metadata": {}}
