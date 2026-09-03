import datetime
import hmac
import os
import io
import csv
from pathlib import Path

import mercadopago
import streamlit as st
import docx
import requests
from dotenv import load_dotenv

from reviewer import evaluate_project

load_dotenv()
st.set_page_config(page_title="Perspecta Salud", page_icon="📘", layout="wide", initial_sidebar_state="collapsed")

PRICE = 4.90
APP_URL = st.secrets.get("APP_URL", os.getenv("APP_URL", "https://revisordetesis.streamlit.app/"))
API_KEY = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY", ""))
MP_TOKEN = st.secrets.get("MP_ACCESS_TOKEN", os.getenv("MP_ACCESS_TOKEN", ""))
# Aceptamos ambas para que no haya problemas si usó la mía o la nueva
ACCESS_CODE = st.secrets.get("ADMIN_PASSWORD", st.secrets.get("INSTITUTIONAL_CODE", "PERSPECTA2026"))
NORMATIVAS = Path(__file__).parent / "normativas"

# CRM Local Database
CRM_FILE = "clientes_crm.csv"
def registrar_operacion(proyecto, nivel, etapa):
    try:
        file_exists = os.path.isfile(CRM_FILE)
        with open(CRM_FILE, mode='a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(["Fecha", "Proyecto", "Nivel", "Etapa"])
            date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            writer.writerow([date_str, proyecto, nivel, etapa])
    except:
        pass

def create_docx(report_text, project_name):
    doc = docx.Document()
    
    # Agregar Encabezado (Header)
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Perspecta Salud - Revisión Académica Inteligente"
    
    # Agregar Pie de Página (Footer)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Contacto: perspectasalud@gmail.com | Web: https://revisordetesis.streamlit.app/"
    
    # Títulos principales en la primera página
    doc.add_heading("Perspecta Salud", 0)
    doc.add_heading("Informe de Auditoría Científica", 1)
    doc.add_heading(f"Proyecto: {project_name}", 2)
    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y')}")
    
    for para in report_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
:root{--navy:#072e68;--cyan:#08b8be;--ink:#112d54;--muted:#63738c;--pale:#f3f9ff;--line:#deebf4}
html{scroll-behavior:smooth}body,[class*=css]{font-family:Inter,sans-serif;color:var(--ink)}
header[data-testid=stHeader]{height:0;background:transparent}#MainMenu,footer,[data-testid=stSidebarCollapsedControl]{display:none}.block-container{max-width:1240px;padding:0 2rem 2rem}.stMainBlockContainer{padding-top:0!important}
h1,h2,h3{color:var(--navy);letter-spacing:-.03em}.anchor{scroll-margin-top:80px}.nav{height:74px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:10;background:#fffffff2;border-bottom:1px solid #edf3f7;margin:0 -2rem;padding:0 2rem}.brand{font-size:22px;font-weight:800;color:var(--navy)}.brand:before{content:'✚';display:inline-grid;place-items:center;width:36px;height:42px;background:var(--navy);color:white;margin-right:10px;clip-path:polygon(15% 8%,85% 0,100% 92%,30% 100%)}.brand span{color:var(--cyan)}.navlinks{display:flex;align-items:center;gap:25px}.nav a{font-size:13px;font-weight:600;text-decoration:none;color:var(--ink)}.nav .cta,.button{background:var(--cyan);color:#fff;padding:13px 19px;border-radius:8px;box-shadow:0 9px 24px #08b8be2c}
.hero{display:grid;grid-template-columns:1fr 1.08fr;align-items:center;gap:38px;min-height:505px}.eyebrow{font-size:12px;color:var(--cyan);font-weight:800;letter-spacing:.11em;text-transform:uppercase}.hero h1{font-size:clamp(42px,5vw,64px);line-height:1.06;margin:12px 0 20px}.hero p,.head p{font-size:17px;line-height:1.65;color:var(--muted)}.actions{display:flex;align-items:center;gap:22px;margin:27px 0}.button{display:inline-block;text-decoration:none!important;font-weight:700}.textlink{color:var(--navy);text-decoration:none;font-weight:700}.trust{display:flex;gap:18px;flex-wrap:wrap;font-size:12px;font-weight:600;color:#415570}.trust span:before{content:'✓';color:var(--cyan);margin-right:6px}.hero-media{height:430px;position:relative;overflow:hidden;border-radius:0 0 0 65px}.hero-media img{width:100%;height:100%;object-fit:cover}.score{position:absolute;right:24px;top:42px;background:#fffffff2;padding:20px;border-radius:16px;width:170px;box-shadow:0 16px 42px #072e6830;text-align:center}.score b{width:88px;height:88px;margin:12px auto;display:grid;place-items:center;border-radius:50%;border:10px solid var(--cyan);font-size:27px;color:var(--navy)}.score small{color:var(--muted)}
.stats{display:grid;grid-template-columns:repeat(4,1fr);background:var(--pale);border-radius:17px;padding:24px;margin:10px 0 68px}.stat{text-align:center;border-right:1px solid var(--line)}.stat:last-child{border:0}.stat strong{font-size:23px;color:var(--navy)}.stat span{display:block;color:var(--muted);font-size:12px;margin-top:4px}.section{padding:68px 0}.head{text-align:center;max-width:680px;margin:0 auto 38px}.head h2{font-size:35px;margin:0 0 10px}.steps{display:grid;grid-template-columns:repeat(4,1fr);gap:24px}.step{text-align:center;padding:10px 18px}.number{width:58px;height:58px;border-radius:50%;background:#e8f9fa;display:grid;place-items:center;margin:auto;color:var(--cyan);font-size:25px;font-weight:800}.step h3{margin:15px 0 7px}.step p{font-size:13px;line-height:1.6;color:var(--muted)}
.soft{background:linear-gradient(#f6fbff,#fff);margin:0 -2rem;padding-left:2rem;padding-right:2rem}.areas{display:grid;grid-template-columns:repeat(5,1fr);gap:13px}.card{border:1px solid var(--line);border-radius:14px;padding:20px 15px;background:#fff;box-shadow:0 8px 25px #102b5210}.card h3{font-size:14px;min-height:36px}.meter{height:7px;background:#e2edf4;border-radius:5px;overflow:hidden}.meter i{display:block;width:78%;height:100%;background:var(--cyan)}.card small{display:block;margin-top:9px;color:var(--muted)}.human{display:grid;grid-template-columns:1fr 1fr;background:#eef8ff;border-radius:20px;overflow:hidden;margin-top:52px}.human img{width:100%;height:315px;object-fit:cover}.quote{padding:45px;display:flex;flex-direction:column;justify-content:center}.quote b{font-size:58px;color:var(--cyan);line-height:.5}.quote p{font-size:20px;line-height:1.6}.quote small{color:var(--muted)}
.report{display:grid;grid-template-columns:1.15fr .85fr;gap:25px;align-items:center}.report-card{border:1px solid var(--line);border-radius:17px;padding:25px;box-shadow:0 16px 40px #102b5214}.tabs{display:flex;gap:25px;border-bottom:1px solid var(--line);padding-bottom:12px;font-size:12px}.tabs b{color:var(--cyan)}.result{display:grid;grid-template-columns:.7fr 1.3fr;gap:25px;padding-top:25px}.donut{width:105px;height:105px;border:14px solid var(--cyan);border-radius:50%;display:grid;place-items:center;margin:auto;font-size:25px;font-weight:800;color:var(--navy)}.bar{font-size:11px;margin:12px 0}.bar i{display:block;height:7px;background:linear-gradient(90deg,var(--cyan) 78%,#e5edf4 78%);border-radius:5px;margin-top:5px}.benefits{display:grid;gap:13px}.benefits div{background:var(--pale);padding:17px;border-radius:11px}.benefits b{display:block;margin-bottom:4px}.benefits span{font-size:13px;color:var(--muted)}
.price{display:grid;grid-template-columns:1.1fr .9fr .65fr;align-items:center;gap:25px;border:1px solid var(--line);border-radius:19px;padding:34px;box-shadow:0 15px 40px #102b5212}.amount{font-size:40px;color:var(--navy);font-weight:800}.checks{line-height:2;color:#41536d}.checks div:before{content:'✓';color:var(--cyan);font-weight:900;margin-right:7px}.security{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}.security article{border:1px solid var(--line);border-radius:13px;padding:22px}.security h3{font-size:16px}.security p,.faq p{font-size:13px;line-height:1.6;color:var(--muted)}.faq{display:grid;grid-template-columns:1fr 1fr;gap:11px}.faq details{border:1px solid var(--line);border-radius:9px;padding:16px}.faq summary{font-weight:600;cursor:pointer}.final{display:flex;align-items:center;justify-content:space-between;background:var(--navy);color:white;border-radius:19px;padding:38px 48px}.final h2{color:white;margin:0 0 7px}.final p{margin:0;color:#d9e7f7}.footer{display:flex;justify-content:space-between;align-items:center;padding:35px 0;color:var(--muted);font-size:12px}.app-head{text-align:center;padding:48px 0 22px}.app-head h1{font-size:42px}.app-head p{color:var(--muted)}div.stButton>button,div.stLinkButton>a{background:var(--cyan);border:0;color:white;border-radius:8px;min-height:48px;font-weight:700}.stDownloadButton>button{background:var(--navy)!important;color:white!important}
@media(max-width:900px){.navlinks a:not(.cta){display:none}.hero,.human,.report,.price{grid-template-columns:1fr}.hero{padding-top:35px}.stats,.steps{grid-template-columns:1fr 1fr}.areas{grid-template-columns:1fr 1fr}.security{grid-template-columns:1fr}.hero-media{border-radius:20px}.faq{grid-template-columns:1fr}.final{align-items:flex-start;gap:20px;flex-direction:column}}
@media(max-width:560px){.block-container{padding:0 1rem 1rem}.nav{margin:0 -1rem;padding:0 1rem}.brand{font-size:16px}.brand:before{width:27px;height:33px}.stats,.steps,.areas{grid-template-columns:1fr}.stat{border:0;border-bottom:1px solid var(--line);padding:10px}.soft{margin:0 -1rem;padding-left:1rem;padding-right:1rem}.hero h1{font-size:40px}.hero-media{height:310px}.score{width:140px;padding:12px}.result{grid-template-columns:1fr}.price{padding:24px}.footer{flex-direction:column;gap:18px}.section{padding:52px 0}}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# Google Analytics Inject
ga_code = """
<!-- Google tag (gtag.js) -->
<script async src="https://www.googletagmanager.com/gtag/js?id=G-J2NW5JD21D"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){dataLayer.push(arguments);}
  gtag('js', new Date());
  gtag('config', 'G-J2NW5JD21D');
</script>
"""
import streamlit.components.v1 as components
components.html(ga_code, height=0, width=0)

def nav():
    st.markdown("""<nav class="nav"><div class="brand">Perspecta <span>Salud</span></div><div class="navlinks"><a href="#inicio">Inicio</a><a href="#como">Cómo funciona</a><a href="#evaluamos">Qué evaluamos</a><a href="#informe">Informe</a><a href="#faq">Preguntas frecuentes</a><a class="cta" href="#empezar">Revisar mi proyecto →</a></div></nav>""", unsafe_allow_html=True)


def verify_payment(payment_id):
    if not MP_TOKEN or not payment_id:
        return False
    try:
        data = mercadopago.SDK(MP_TOKEN).payment().get(str(payment_id)).get("response", {})
        return data.get("status") == "approved" and data.get("currency_id") == "PEN" and float(data.get("transaction_amount", 0)) >= PRICE
    except Exception:
        return False


def payment_link():
    data = {"items":[{"title":"Revisión académica - Perspecta Salud","quantity":1,"unit_price":PRICE,"currency_id":"PEN"}],"back_urls":{"success":APP_URL,"failure":APP_URL,"pending":APP_URL},"auto_return":"approved","statement_descriptor":"PERSPECTA SALUD"}
    return mercadopago.SDK(MP_TOKEN).preference().create(data)["response"]["init_point"]


def landing():
    nav()
    st.markdown(f"""
    <main><section id="inicio" class="hero anchor"><div><div class="eyebrow">Revisión académica inteligente</div><h1>Mejora tu proyecto de investigación con una revisión inteligente</h1><p>Recibe observaciones claras, criterios de cumplimiento y recomendaciones para fortalecer tu trabajo académico.</p><div class="actions"><a class="button" href="#empezar">Revisar mi proyecto →</a><a class="textlink" href="#informe">Ver informe de ejemplo ›</a></div><div class="trust"><span>Basada en rúbrica</span><span>Informe descargable en Word</span><span>Documentos confidenciales</span></div></div><div class="hero-media"><img src="https://images.unsplash.com/photo-1541339907198-e08756dedf3f?auto=format&fit=crop&w=1200&q=86" alt="Estudiante trabajando en su investigación"><div class="score"><small>Resultado general</small><b>78</b><small>Buen avance</small></div></div></section>
    <section class="stats"><div class="stat"><strong>5</strong><span>áreas evaluadas</span></div><div class="stat"><strong>⏱</strong><span>Informe en minutos</span></div><div class="stat"><strong>24/7</strong><span>Siempre disponible</span></div><div class="stat"><strong>S/ {PRICE:.2f}</strong><span>Precio introductorio</span></div></section>
    <section id="como" class="section anchor"><div class="head"><h2>¿Cómo funciona?</h2><p>Cuatro pasos sencillos para saber qué debes mejorar.</p></div><div class="steps"><article class="step"><div class="number">1</div><h3>Selecciona</h3><p>Elige la revisión que se adapta a tu proyecto.</p></article><article class="step"><div class="number">2</div><h3>Paga</h3><p>Realiza el pago en línea de forma segura.</p></article><article class="step"><div class="number">3</div><h3>Carga</h3><p>Sube tu proyecto y anexos relevantes.</p></article><article class="step"><div class="number">4</div><h3>Descarga</h3><p>Recibe observaciones y recomendaciones en formato Word profesional.</p></article></div></section>
    <section id="evaluamos" class="section soft anchor"><div class="head"><h2>Descubre qué debes mejorar</h2><p>Evaluamos tu proyecto con criterios académicos en cinco áreas clave.</p></div><div class="areas"><article class="card"><h3>⚛ Coherencia científica</h3><div class="meter"><i style="width:82%"></i></div><small>Problema, objetivos e hipótesis</small></article><article class="card"><h3>▣ Metodología</h3><div class="meter"><i style="width:74%"></i></div><small>Diseño, población e instrumentos</small></article><article class="card"><h3>▤ Fuentes bibliográficas</h3><div class="meter"><i style="width:76%"></i></div><small>Actualidad y pertinencia</small></article><article class="card"><h3>❞ Estilo de citación</h3><div class="meter"><i style="width:80%"></i></div><small>Aplicación del estilo solicitado</small></article><article class="card"><h3>♢ Normativa institucional</h3><div class="meter"><i></i></div><small>Estructura y requisitos formales</small></article></div><div class="human"><img src="https://images.unsplash.com/photo-1523240795612-9a054b0db644?auto=format&fit=crop&w=1200&q=86" alt="Estudiantes revisando un proyecto"><div class="quote"><b>“</b><p>Observaciones claras para que sepas por dónde empezar y qué aspectos requieren mayor atención.</p><small>Una revisión diseñada para estudiantes y asesores</small></div></div></section>
    <section id="informe" class="section anchor"><div class="head"><h2>Conoce el informe que recibirás</h2><p>Resultados organizados para convertir las observaciones en acciones concretas.</p></div><div class="report"><div class="report-card"><div class="tabs"><b>Resumen</b><span>Observaciones</span><span>Recomendaciones</span></div><div class="result"><div><div class="donut">78</div><p style="text-align:center;color:var(--cyan);font-weight:700">Buen avance</p></div><div><div class="bar">Coherencia científica<i></i></div><div class="bar">Metodología<i style="background:linear-gradient(90deg,var(--cyan) 74%,#e5edf4 74%)"></i></div><div class="bar">Fuentes bibliográficas<i style="background:linear-gradient(90deg,var(--cyan) 76%,#e5edf4 76%)"></i></div><div class="bar">Estilo de citación<i style="background:linear-gradient(90deg,var(--cyan) 80%,#e5edf4 80%)"></i></div></div></div></div><div class="benefits"><div><b>✓ Observaciones claras</b><span>Identifica los puntos que necesitan atención.</span></div><div><b>✓ Recomendaciones aplicables</b><span>Conoce qué acción realizar en cada criterio.</span></div><div><b>🔒 Documento confidencial</b><span>Solo tú tendrás acceso al informe.</span></div></div></div></section>
    <section id="empezar" class="section anchor"><div class="price"><div><h3>Revisión introductoria</h3><p style="color:var(--muted)">Ideal para proyectos de grado y posgrado.</p><div class="amount">S/ {PRICE:.2f}</div><small>Precio promocional por tiempo limitado.</small></div><div class="checks"><div>Evaluación basada en rúbrica</div><div>Observaciones y recomendaciones</div><div>Informe descargable en Word (.docx)</div></div><div><a class="button" href="#activar">Revisar ahora →</a></div></div></section>
    <section class="section"><div class="head"><h2>Tu investigación merece confidencialidad</h2></div><div class="security"><article><h3>🛡 Archivos protegidos</h3><p>Se procesan únicamente para generar la evaluación.</p></article><article><h3>🧠 Uso responsable de IA</h3><p>La IA asiste; la decisión académica final corresponde a expertos.</p></article><article><h3>◷ Eliminación programada</h3><p>Los archivos temporales no se conservan indefinidamente.</p></article></div></section>
    <section id="faq" class="section anchor"><div class="head"><h2>Preguntas frecuentes</h2></div><div class="faq"><details><summary>¿Qué proyectos pueden revisar?</summary><p>Trabajos de pregrado y posgrado compatibles con la rúbrica habilitada.</p></details><details><summary>¿Mis documentos están seguros?</summary><p>Se usan únicamente para generar tu evaluación.</p></details><details><summary>¿Cuánto demora?</summary><p>Generalmente pocos minutos, según la extensión.</p></details><details><summary>¿Reemplaza a mi asesor?</summary><p>No. Complementa la orientación académica profesional.</p></details></div></section>
    <section class="final"><div><h2>¿Listo para mejorar tu proyecto?</h2><p>Identifica oportunidades de mejora y lleva tu investigación al siguiente nivel.</p></div><a class="button" href="#activar">Revisar mi proyecto →</a></section>
    """, unsafe_allow_html=True)
    
    # Bloque para Iniciar / FlagCounter / Formulario
    _, center, _ = st.columns([1, 1.25, 1])
    with center:
        if st.button("Revisar mi proyecto por S/ 4.90", type="primary", use_container_width=True):
            st.session_state.page = "review"
            st.rerun()
            
        st.divider()
        st.markdown("<div style='text-align: center'><h3>✉️ Contáctanos</h3><p>¿Dudas? ¿Asesoría Estadística Avanzada? Escríbenos.</p></div>", unsafe_allow_html=True)
        with st.form("contact_form_landing"):
            u_email = st.text_input("Tu correo electrónico:")
            u_msg = st.text_area("Mensaje:")
            if st.form_submit_button("Enviar Mensaje", use_container_width=True):
                contact_email = st.secrets.get("CONTACT_EMAIL", "")
                if contact_email and u_email:
                    try:
                        requests.post(f"https://formsubmit.co/ajax/{contact_email}", json={"Email": u_email, "Mensaje": u_msg})
                        st.success("¡Mensaje enviado con éxito!")
                    except:
                        st.error("Error enviando el mensaje.")

        st.markdown("<div style='text-align: center; margin-top: 30px;'><p>Estadísticas de la plataforma:</p></div>", unsafe_allow_html=True)
        html_code = """<div style="text-align: center;"><a href="https://info.flagcounter.com/OA6D"><img src="https://s01.flagcounter.com/countxl/OA6D/bg_FFFFFF/txt_000000/border_CCCCCC/columns_2/maxflags_4/viewers_Usuarios/labels_0/pageviews_1/flags_0/percent_0/" alt="Flag Counter" border="0"></a></div>"""
        components.html(html_code, height=180)

    st.markdown('<footer class="footer"><div class="brand">Perspecta <span>Salud</span></div><div>Términos · Privacidad · Soporte · Contacto</div><div>© 2026 Perspecta Salud</div></footer></main><div id="activar" class="anchor"></div>', unsafe_allow_html=True)


def review():
    nav()
    st.markdown('<div class="app-head"><div class="eyebrow">Revisión académica</div><h1>Revisa tu proyecto</h1><p>Completa el proceso y recibe un informe organizado.</p></div>', unsafe_allow_html=True)
    paid = bool(st.session_state.get("paid"))
    
    with st.container(border=True):
        if not paid:
            st.subheader("1. Activa tu revisión")
            st.write(f"Revisión introductoria: **S/ {PRICE:.2f}**. El acceso se habilita al comprobar el pago.")
            if MP_TOKEN:
                try: st.link_button("Pagar de forma segura", payment_link(), type="primary", use_container_width=True)
                except Exception: st.error("No pudimos crear el enlace de pago. Inténtalo nuevamente.")
            else: st.warning("El sistema de pagos está temporalmente en mantenimiento.")
            
            with st.expander("Acceso institucional (Administrador)"):
                code = st.text_input("Código institucional", type="password")
                if st.button("Validar código", use_container_width=True):
                    if ACCESS_CODE and code.strip() == str(ACCESS_CODE).strip(): 
                        st.session_state.paid=True
                        st.session_state.is_admin=True
                        st.rerun()
                    else: st.error("El código no es válido.")
            if st.button("← Volver al inicio"): st.session_state.page="home"; st.rerun()
            return

        # ==================== APLICACIÓN DESBLOQUEADA ====================
        st.success("Acceso confirmado. Ya puedes cargar tu proyecto.")
        
        # Panel Secreto del Administrador
        if st.session_state.get("is_admin"):
            st.write("📊 **Panel Secreto de Administrador**")
            if os.path.isfile(CRM_FILE):
                with open(CRM_FILE, "rb") as f:
                    st.download_button("📥 Descargar Base de Datos (CSV)", data=f, file_name="CRM_PerspectaSalud.csv", mime="text/csv")
            else:
                st.info("No hay proyectos procesados aún.")
            st.divider()
        
        # EL CANDADO DE 1 SOLO USO: Si ya hay reporte, ocultamos el formulario
        if st.session_state.get("report"):
            st.warning("⚠️ **IMPORTANTE:** Descarga tu reporte en formato Word ahora. Si recargas o cierras esta página, tu sesión de pago expirará y tendrás que volver a pagar.")
            st.subheader("Resultado de la revisión")
            st.write(st.session_state.report)
            
            # Generar Word File
            date=datetime.datetime.now().strftime("%d%m%Y")
            safe="".join(c for c in st.session_state.project if c.isalnum() or c in "-_ ").strip()
            
            docx_data = create_docx(st.session_state.report, st.session_state.project)
            st.download_button(
                label="📥 Descargar Informe Completo en Word (.docx)", 
                data=docx_data, 
                file_name=f"Auditoria_PerspectaSalud_{safe}_{date}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                use_container_width=True
            )
            st.info("Para auditar un nuevo proyecto, debes recargar la página (F5) e iniciar una nueva sesión.")
            return

        # EL FORMULARIO
        st.subheader("2. Datos del proyecto")
        name = st.text_input("Nombre del autor o proyecto", placeholder="Ej. Apellido – título breve")
        nivel = st.selectbox("Nivel académico", ["Pregrado", "Maestría", "Doctorado"])
        etapa = st.selectbox("Etapa del trabajo", ["Proyecto de investigación", "Tesis en desarrollo", "Tesis final"])
        
        st.write("⚠️ **Límites:** Máximo 3 archivos. Peso total máximo 200 MB.")
        files = st.file_uploader("Carga los documentos", type=["pdf","docx","txt","xlsx","csv"], accept_multiple_files=True)
        consent = st.checkbox("Confirmo que tengo autorización para procesar los documentos y acepto las condiciones de privacidad.")
        
        if st.button("Iniciar revisión", type="primary", use_container_width=True):
            if not API_KEY: st.error("El motor de evaluación no está configurado.")
            elif not name.strip(): st.error("Ingresa el nombre del proyecto.")
            elif not files: st.error("Carga al menos un documento.")
            elif len(files) > 3: st.error("Por favor, sube un máximo de 3 archivos consolidados.")
            elif not consent: st.error("Debes aceptar las condiciones de privacidad.")
            elif any(f.size > 200*1024*1024 for f in files): st.error("Cada archivo debe pesar como máximo 200 MB.")
            else:
                import time
                progress_bar = st.progress(0, text="Iniciando protocolos de lectura...")
                time.sleep(0.5)
                progress_bar.progress(30, text="Extrayendo texto y analizando metodología estadística (SPSS)...")
                time.sleep(1)
                progress_bar.progress(60, text="Verificando citas, referencias y normativa institucional...")
                time.sleep(1)
                progress_bar.progress(85, text="Cruzando datos con el panel de expertos de IA...")

                with st.status("Generando reporte de nivel doctoral...", expanded=True) as status:
                    result = evaluate_project(files, str(NORMATIVAS), API_KEY)
                    status.update(label="Auditoría Completada", state="complete", expanded=False)
                
                progress_bar.progress(100, text="¡Auditoría completada!")
                time.sleep(0.5)
                progress_bar.empty()

                if isinstance(result, str): st.error(result)
                else: 
                    # Registrar en CRM
                    registrar_operacion(name, nivel, etapa)
                    
                    st.session_state.report=result.get("report", "No fue posible recuperar el informe.")
                    st.session_state.project=name
                    st.rerun() # Recarga para ocultar el formulario

payment_id = st.query_params.get("payment_id", "")
status = st.query_params.get("status", "")

if status == "approved" and payment_id and verify_payment(payment_id):
    st.session_state.paid=True
    st.session_state.page="review"
    st.query_params.clear()

st.session_state.setdefault("page", "home")
review() if st.session_state.page == "review" else landing()
